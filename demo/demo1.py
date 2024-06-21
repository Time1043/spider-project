import csv
import os

from docx import Document
from docx.oxml.ns import qn
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait


def get_element_text(parent_element, xpath):
    try:
        element = parent_element.find_element(by=By.XPATH, value=xpath)
        return element.text if element else ''
    except Exception:
        print("====================================================================================1")
        return ''


def get_article_single(href, base_folder):
    # url = "https://mp.weixin.qq.com/mp/appmsg/show?__biz=MjM5MzEwNzYwNQ==&appmsgid=10000020&itemidx=1&sign=258dc8e172087646c9e1998334e27089&scene=21#wechat_redirect"
    # url = "https://mp.weixin.qq.com/mp/appmsg/show?__biz=MjM5MzEwNzYwNQ==&appmsgid=10000023&itemidx=1&sign=fc1391e04a96b58a980ce494c1d6a0b4&scene=21#wechat_redirect"

    chrome_options = Options()
    chrome_options.add_argument("--headless")
    wd = webdriver.Chrome(options=chrome_options)
    wait = WebDriverWait(wd, 10)
    wd.get(href)

    try:
        article = wd.find_element(by=By.XPATH, value='//div[@id="img-content"]')

        title = get_element_text(article, './/h1[@id="activity-name"]')
        name_title = get_element_text(article, './/div[@id="meta_content"]')
        content = get_element_text(article, './/div[@id="js_content"]')
        # name1 = get_element_text(article, './/div[@id="meta_content"]/span[@class="rich_media_meta_text"]')
        # name2 = get_element_text(article, './/div[@id="meta_content"]/span[@class="rich_media_meta_nickname"]')
        # time = get_element_text(article, './/div[@id="meta_content"]/span[@id="meta_content_hide_info"]')

        title = title.replace('\n', '').replace('\r', '') \
            .replace(' ', '_').replace('/', '_') \
            .replace('?', '_').replace('"', '_').replace(':', '_') \
            .replace('|', '_')

        # csv
        with open(csv_file, mode='a', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow([title, name_title, content])

        # word
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        doc.add_heading(str(title), level=1)
        doc.add_paragraph(str(name_title))
        doc.add_paragraph(str(content))

        file_name = f"crawl/{base_folder}/{title}.docx"
        doc.save(file_name)

    except Exception as e:
        print("====================================================================================2")
        print(e)


def get_article_list(list_href, base_folder):
    # url = "https://mp.weixin.qq.com/s?__biz=MjM5MzEwNzYwNQ==&mid=2650023025&idx=1&sn=388c240350a33e3f167d3bb31a3bba84&chksm=be9c9f8b89eb169d7a0d298d4d7b1f90453cbbd4dd90e5201c8ad972170e36d5257f76ed1f6e&scene=178&cur_album_id=2774720597133230081#rd"

    chrome_options = Options()
    chrome_options.add_argument("--headless")
    wd = webdriver.Chrome(options=chrome_options)
    wait = WebDriverWait(wd, 10)
    wd.get(list_href)

    articles_div = wd.find_element(by=By.XPATH, value='//div[@id="js_content"]')
    articles_a = articles_div.find_elements(by=By.XPATH, value='.//p/a')

    hrefs = []
    for item in articles_a:
        hrefs.append(item.get_attribute('href'))

    num = 1
    for href in hrefs:
        print(num)
        get_article_single(href, base_folder)
        num += 1


def get_article_list_list():
    url = "https://mp.weixin.qq.com/mp/appmsgalbum?__biz=MjM5MzEwNzYwNQ==&action=getalbum&album_id=2774720597133230081&scene=126&sessionid=-1900623423&uin=&key=&devicetype=Windows+11+x64&version=6309092b&lang=zh_CN&ascene=0"

    chrome_options = Options()
    chrome_options.add_argument("--headless")
    wd = webdriver.Chrome(options=chrome_options)
    wait = WebDriverWait(wd, 10)
    wd.get(url)

    articles_list_ul = wd.find_element(by=By.XPATH, value='//ul[@class="album__list js_album_list"]')
    articles_list_li = articles_list_ul.find_elements(by=By.XPATH, value='.//li')

    list_hrefs = []
    for item in articles_list_li:
        list_hrefs.append(item.get_attribute('data-link'))

    base_folder = 1
    for list_href in list_hrefs:
        print(f"%%%%%{base_folder}%%%%%")
    os.mkdir(f"crawl/{base_folder}")
    get_article_list(list_href, base_folder)
    base_folder += 1

csv_file = 'crawl/output.csv'


try:
    with open(csv_file, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(['标题', '名称和时间', '内容'])
except Exception as e:
    print(f"Error opening or creating CSV file: {e}")

get_article_list_list()
